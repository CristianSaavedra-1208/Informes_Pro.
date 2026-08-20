import os
import sys
import pandas as pd
from io import BytesIO
from docx import Document
from docxtpl import DocxTemplate
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# Asegurar importabilidad de src
ROOT_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
sys.path.append(ROOT_DIR)

from src.models.database import SessionLocal
from src.models.consolidacion import ConsolidationGroup
from src.models.trial_balance_db import TrialBalanceDB
from src.models.pl_cubo_db import PlCuboDB

from docx.shared import Inches, Pt, RGBColor

def set_cell_bg(cell, fill_color):
    """Inyecta XML shading para cambiar el color de fondo de una celda de Word."""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill_color))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=50, bottom=50, left=100, right=100):
    """Ajusta los márgenes internos (padding) de la celda de Word en dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_cell_borders(cell, top=None, bottom=None):
    """Inyecta bordes OXML a una celda de Word (útil para filas de totales/subtotales contables)."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders_xml = f'<w:tcBorders {nsdecls("w")}>'
    if top:
        borders_xml += f'<w:top w:val="{top.get("val", "single")}" w:sz="{top.get("sz", "6")}" w:space="0" w:color="{top.get("color", "000000")}"/>'
    else:
        borders_xml += f'<w:top w:val="none"/>'
    if bottom:
        borders_xml += f'<w:bottom w:val="{bottom.get("val", "single")}" w:sz="{bottom.get("sz", "6")}" w:space="0" w:color="{bottom.get("color", "000000")}"/>'
    else:
        borders_xml += f'<w:bottom w:val="none"/>'
    borders_xml += f'<w:left w:val="none"/><w:right w:val="none"/>'
    borders_xml += '</w:tcBorders>'
    tcPr.append(parse_xml(borders_xml))

def apply_table_column_widths(table, num_cols, total_width_inches=6.8):
    """Calcula y aplica anchos de columna óptimos para informes financieros."""
    table.autofit = False
    
    if num_cols <= 1:
        col_widths = [total_width_inches]
    elif num_cols == 2:
        col_widths = [total_width_inches * 0.70, total_width_inches * 0.30]
    elif num_cols == 3:
        col_widths = [total_width_inches * 0.64, total_width_inches * 0.18, total_width_inches * 0.18]
    elif num_cols == 4:
        col_widths = [total_width_inches * 0.58, total_width_inches * 0.10, total_width_inches * 0.16, total_width_inches * 0.16]
    else:
        col0 = max(3.2, total_width_inches - (1.1 * (num_cols - 1)))
        rem = max(0.8, (total_width_inches - col0) / (num_cols - 1))
        col_widths = [col0] + [rem] * (num_cols - 1)

    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        for j, cell in enumerate(row.cells):
            if j < len(col_widths):
                cell.width = Inches(col_widths[j])

    if len(table.rows) > 0:
        hdr_trPr = table.rows[0]._tr.get_or_add_trPr()
        hdr_trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

def move_element_before(target_p, new_element):
    """Mueve un elemento XML de Word (párrafo o tabla) justo antes del párrafo objetivo."""
    parent = target_p._p.getparent()
    idx = parent.index(target_p._p)
    parent.insert(idx, new_element)

def populate_docx_table(table, df, numeric_cols):
    """Rellena y da estilo corporativo a una tabla de Word en base a un DataFrame."""
    table.style = 'Normal Table'
    num_cols = len(df.columns)
    
    # Fila de cabecera
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = str(col_name)
        set_cell_bg(hdr_cells[i], "1F4E78") # Azul corporativo
        set_cell_margins(hdr_cells[i], top=70, bottom=70, left=100, right=100)
        for p in hdr_cells[i].paragraphs:
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.0
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255) # Blanco
                run.font.name = 'Arial'
                run.font.size = Pt(9)
                
    # Filas de datos
    for row_i, (_, row_vals) in enumerate(df.iterrows()):
        row_cells = table.rows[row_i + 1].cells
        first_cell_str = str(row_vals.iloc[0]).strip()
        first_cell_clean = first_cell_str.lower()
        
        is_grand_total = any(k in first_cell_clean for k in ["total activos", "total patrimonio y pasivos", "patrimonio total", "ganancia (pérdida) del ejercicio", "resultado del ejercicio"])
        is_subtotal = not is_grand_total and any(k in first_cell_clean for k in ["total", "sub total", "saldo final", "ganancia bruta", "resultado antes de", "procedente de operaciones"])
        is_category_header = not is_subtotal and not is_grand_total and (
            first_cell_str in ["Activos", "Activos corrientes", "Activos no corrientes", "Patrimonio y pasivos", "Pasivos corrientes", "Pasivos no corrientes", "Patrimonio"] or
            all(pd.isna(row_vals.iloc[col_i]) or str(row_vals.iloc[col_i]).strip() == "" for col_i in range(1, num_cols))
        )

        if is_grand_total:
            bg_color = "EBF2FA"
        elif is_subtotal:
            bg_color = "F4F7FA"
        elif is_category_header:
            bg_color = "F9FAFB"
        else:
            is_even = (row_i % 2 == 0)
            bg_color = "FFFFFF" if is_even else "F9FAFB"

        for col_idx, col_name in enumerate(df.columns):
            val = row_vals[col_name]
            
            if pd.isna(val) or val == "None" or str(val) == "nan":
                text_val = ""
            else:
                if isinstance(val, (int, float)):
                    try:
                        if val == 0:
                            text_val = "-"
                        elif val < 0:
                            text_val = f"({abs(val):,.0f})".replace(",", ".")
                        else:
                            text_val = f"{val:,.0f}".replace(",", ".")
                    except:
                        text_val = str(val)
                else:
                    text_val = str(val)
                    
            cell = row_cells[col_idx]
            cell.text = text_val
            set_cell_bg(cell, bg_color)
            set_cell_margins(cell, top=40, bottom=40, left=100, right=100)
            
            if is_grand_total:
                set_cell_borders(cell, top={'val': 'single', 'sz': '6'}, bottom={'val': 'double', 'sz': '12'})
            elif is_subtotal:
                set_cell_borders(cell, top={'val': 'single', 'sz': '6'}, bottom={'val': 'single', 'sz': '6'})

            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1.5)
                p.paragraph_format.space_after = Pt(1.5)
                p.paragraph_format.line_spacing = 1.0
                
                if col_idx in numeric_cols or isinstance(val, (int, float)) or col_idx > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(8.5)
                    if is_grand_total or is_subtotal or is_category_header:
                        run.font.bold = True
                        if is_grand_total or is_category_header:
                            run.font.color.rgb = RGBColor(15, 45, 80)

    apply_table_column_widths(table, num_cols, total_width_inches=6.8)

def populate_notes_docx_table(table, chunk_df):
    """Rellena y da estilo a una tabla de notas en base a un DataFrame de openpyxl/slicing."""
    table.style = 'Normal Table'
    num_cols = len(chunk_df.columns)
    numeric_cols = chunk_df.select_dtypes(include=['number']).columns.tolist()
    
    in_header = True
    for row_i, (_, row) in enumerate(chunk_df.iterrows()):
        row_cells = table.rows[row_i].cells
        
        has_data_numbers = False
        for col_idx, val in enumerate(row):
            if pd.notna(val) and isinstance(val, (int, float)):
                if not (val > 2000 and val < 2100):
                    has_data_numbers = True
                    break
        if has_data_numbers:
            in_header = False
            
        first_cell_val = str(row.iloc[0]).lower().strip()
        is_grand_total = any(x in first_cell_val for x in ["total general", "total final", "saldo final", "total patrimonio", "total activos"])
        is_subtotal = not is_grand_total and any(x in first_cell_val for x in ["total", "sub total", "subtotal", "ganancia bruta"])
        
        if in_header:
            bg_color = "1F4E78"
        elif is_grand_total:
            bg_color = "EBF2FA"
        elif is_subtotal:
            bg_color = "F4F7FA"
        else:
            is_even = (row_i % 2 == 0)
            bg_color = "FFFFFF" if is_even else "F9FAFB"
            
        for col_idx, col_name in enumerate(chunk_df.columns):
            val = row[col_name]
            if pd.isna(val) or val == "None" or str(val) == "nan" or str(val).strip() == "":
                text_val = ""
            else:
                if isinstance(val, (int, float)):
                    try:
                        if val == 0:
                            text_val = "-"
                        elif val < 0:
                            text_val = f"({abs(val):,.0f})".replace(",", ".")
                        else:
                            text_val = f"{val:,.0f}".replace(",", ".")
                    except:
                        text_val = str(val)
                else:
                    text_val = str(val)
                    
            cell = row_cells[col_idx]
            cell.text = text_val
            set_cell_bg(cell, bg_color)
            set_cell_margins(cell, top=40, bottom=40, left=100, right=100)
            
            if is_grand_total:
                set_cell_borders(cell, top={'val': 'single', 'sz': '6'}, bottom={'val': 'double', 'sz': '12'})
            elif is_subtotal:
                set_cell_borders(cell, top={'val': 'single', 'sz': '6'}, bottom={'val': 'single', 'sz': '6'})
                
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(1.5)
                p.paragraph_format.space_after = Pt(1.5)
                p.paragraph_format.line_spacing = 1.0
                
                if col_idx > 0 or col_idx in numeric_cols or isinstance(val, (int, float)):
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(8.5)
                    if in_header:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    elif is_grand_total or is_subtotal:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(15, 45, 80)
                        
    apply_table_column_widths(table, num_cols, total_width_inches=6.8)

def load_mappings_for_entity(empresa):
    """Carga los mapeos de balance y PL para una entidad."""
    empresa_path = os.path.join("data", "empresas", empresa)
    
    map_balance_df = None
    map_pl_df = None
    
    # 1. Intentar cargar desde los archivos individuales específicos (Source of Truth)
    p_bal = os.path.join(empresa_path, "map_balance.xlsx")
    if os.path.exists(p_bal):
        try:
            map_balance_df = pd.read_excel(p_bal)
        except:
            pass
            
    p_pl = os.path.join(empresa_path, "map_pl.xlsx")
    if os.path.exists(p_pl):
        try:
            map_pl_df = pd.read_excel(p_pl)
        except:
            pass
            
    # 2. Si no existen, intentar cargar desde la plantilla de notas global
    if map_balance_df is None or map_pl_df is None:
        template_nota = "Plantilla de notas_v1.xlsx"
        if os.path.exists(template_nota):
            if map_balance_df is None:
                try:
                    map_balance_df = pd.read_excel(template_nota, sheet_name="Mapeo Balance")
                except:
                    pass
            if map_pl_df is None:
                try:
                    map_pl_df = pd.read_excel(template_nota, sheet_name="Mapeo Ctas P&L Cubo")
                except:
                    pass
                    
    return map_balance_df, map_pl_df

class WordTemplateEngine:
    def __init__(self, template_bytes_io):
        self.template_bytes_io = template_bytes_io

    def _extract_cf_preview(self, result_bytes, col_actual, col_comp):
        result_bytes.seek(0)
        import openpyxl
        import datetime
        import re
        wb_check = openpyxl.load_workbook(result_bytes, data_only=True)
        ws_check = wb_check.active
        
        name_col_idx = 1
        for col in range(1, 10):
            for row in range(1, 15):
                val = ws_check.cell(row=row, column=col).value
                if val and str(val).strip().lower() in ["concepto", "descripcion", "detalle", "flujos", "origen/aplicacion"]:
                    name_col_idx = col
                    break
                    
        date_cols = []
        for col in range(1, ws_check.max_column + 1):
            if col == name_col_idx:
                continue
            for row in range(1, 10):
                val = ws_check.cell(row=row, column=col).value
                if val is not None:
                    is_date = (
                        isinstance(val, (datetime.datetime, datetime.date)) or
                        (isinstance(val, str) and re.search(r'20\d{2}', val))
                    )
                    if is_date:
                        date_cols.append(col)
                        break
        date_cols = sorted(list(set(date_cols)))
        
        val25_col_idx = 3
        val24_col_idx = 4
        if len(date_cols) >= 2:
            val25_col_idx = date_cols[0]
            val24_col_idx = date_cols[1]
            
        result_bytes.seek(0)
        df_raw = pd.read_excel(result_bytes)
        cols_to_keep = [name_col_idx - 1, val25_col_idx - 1, val24_col_idx - 1]
        preview_df = df_raw.iloc[:, cols_to_keep].copy()
        preview_df.columns = ["Descripción", col_actual, col_comp]
        preview_df = preview_df.dropna(how='all', subset=["Descripción"]).reset_index(drop=True)
        
        clean_rows = []
        for _, r in preview_df.iterrows():
            desc = str(r["Descripción"]).strip().lower()
            if any(t in desc for t in ["estados separados", "por los ejercicios terminados"]):
                continue
            clean_rows.append(r)
        preview_df = pd.DataFrame(clean_rows).reset_index(drop=True) if clean_rows else preview_df

        return preview_df

    def _extract_balance_preview(self, result_bytes, col_actual, col_comp):
        result_bytes.seek(0)
        import openpyxl
        import datetime
        import re
        wb_check = openpyxl.load_workbook(result_bytes, data_only=True)
        ws_check = wb_check.active
        
        # Buscar dinámicamente las columnas en el excel generado
        from src.core.excel_utils import detect_balance_columns
        name_col_idx, nota_col_idx, val25_col_idx, val24_col_idx = detect_balance_columns(ws_check)
            
        result_bytes.seek(0)
        df_raw = pd.read_excel(result_bytes)
        
        if nota_col_idx is not None:
            cols_to_keep = [name_col_idx - 1, nota_col_idx - 1, val25_col_idx - 1, val24_col_idx - 1]
            preview_df = df_raw.iloc[:, cols_to_keep].copy()
            preview_df.columns = ["Clasificación", "Nota", col_actual, col_comp]
            
            # Clean the 'Nota' column
            def clean_nota(val):
                if pd.isna(val) or str(val).strip() == "" or str(val).lower() == "nan" or val == 0:
                    return ""
                try:
                    float_val = float(val)
                    if float_val.is_integer():
                        return str(int(float_val))
                    return str(float_val)
                except:
                    return str(val).strip()
            
            preview_df["Nota"] = preview_df["Nota"].apply(clean_nota)
        else:
            cols_to_keep = [name_col_idx - 1, val25_col_idx - 1, val24_col_idx - 1]
            preview_df = df_raw.iloc[:, cols_to_keep].copy()
            preview_df.columns = ["Clasificación", col_actual, col_comp]
            
        preview_df = preview_df.dropna(subset=["Clasificación"]).reset_index(drop=True)
        return preview_df

    def _extract_er_preview(self, result_bytes, col_actual, col_comp):
        result_bytes.seek(0)
        import openpyxl
        import datetime
        import re
        wb_check = openpyxl.load_workbook(result_bytes, data_only=True)
        ws_check = wb_check.active
        
        name_col_idx = 1
        for col in range(1, 10):
            for row in range(1, 15):
                val = ws_check.cell(row=row, column=col).value
                if val and str(val).strip().lower() in ["clasificacion", "descripcion", "detalle", "concepto", "ingresos"]:
                    name_col_idx = col
                    break
                    
        date_cols = []
        for col in range(1, ws_check.max_column + 1):
            if col == name_col_idx:
                continue
            for row in range(1, 10):
                val = ws_check.cell(row=row, column=col).value
                if val is not None:
                    is_date = (
                        isinstance(val, (datetime.datetime, datetime.date)) or
                        (isinstance(val, str) and re.search(r'20\d{2}', val))
                    )
                    if is_date:
                        date_cols.append(col)
                        break
        date_cols = sorted(list(set(date_cols)))
        
        val25_col_idx = 2
        val24_col_idx = 3
        if len(date_cols) >= 2:
            val25_col_idx = date_cols[0]
            val24_col_idx = date_cols[1]
            
        result_bytes.seek(0)
        df_raw = pd.read_excel(result_bytes)
        cols_to_keep = [name_col_idx - 1, val25_col_idx - 1, val24_col_idx - 1]
        preview_df = df_raw.iloc[:, cols_to_keep].copy()
        preview_df.columns = ["Clasificación", col_actual, col_comp]
        preview_df = preview_df.dropna(how='all', subset=["Clasificación"]).reset_index(drop=True)
        return preview_df

    def _extract_ori_preview(self, result_bytes, col_actual, col_comp):
        result_bytes.seek(0)
        import openpyxl
        import datetime
        import re
        wb_check = openpyxl.load_workbook(result_bytes, data_only=True)
        ws_check = wb_check.active
        
        name_col_idx = 1
        for col in range(1, 10):
            for row in range(1, 15):
                val = ws_check.cell(row=row, column=col).value
                if val and str(val).strip().lower() in ["detalle", "concepto", "clasificacion", "descripcion"]:
                    name_col_idx = col
                    break
                    
        date_cols = []
        for col in range(1, ws_check.max_column + 1):
            if col == name_col_idx:
                continue
            for row in range(1, 10):
                val = ws_check.cell(row=row, column=col).value
                if val is not None:
                    is_date = (
                        isinstance(val, (datetime.datetime, datetime.date)) or
                        (isinstance(val, str) and re.search(r'20\d{2}', val))
                    )
                    if is_date:
                        date_cols.append(col)
                        break
        date_cols = sorted(list(set(date_cols)))
        
        val25_col_idx = 2
        val24_col_idx = 3
        if len(date_cols) >= 2:
            val25_col_idx = date_cols[0]
            val24_col_idx = date_cols[1]
            
        result_bytes.seek(0)
        from src.core.excel_utils import detect_general_skiprows
        ori_skip = detect_general_skiprows(result_bytes)
        result_bytes.seek(0)
        df_raw = pd.read_excel(result_bytes, skiprows=ori_skip)
        cols_to_keep = [name_col_idx - 1, val25_col_idx - 1, val24_col_idx - 1]
        preview_df = df_raw.iloc[:, cols_to_keep].copy()
        preview_df.columns = ["Detalle", col_actual, col_comp]
        preview_df = preview_df.dropna(how='all', subset=["Detalle"]).reset_index(drop=True)
        return preview_df


    def process(self, empresa_activa, periodo_actual, periodo_comp, scale_factor=1.0, use_ifrs_auto=True):
        doc = DocxTemplate(self.template_bytes_io)
        
        is_consolidated = empresa_activa.startswith("[GRUPO]")
        empresa_path = os.path.join("data", "empresas", empresa_activa)
        
        # 1. Cargar mapeos de balance/PL
        map_balance_df, map_pl_df = load_mappings_for_entity(empresa_activa)
        
        # 2. Calcular variables de texto
        periodo_actual_str = str(periodo_actual)
        periodo_comp_str = str(periodo_comp)
        
        # Calcular periodos comparativos específicos según reglas IFRS o manuales
        periodo_comp_balance = periodo_comp
        periodo_comp_resultados = periodo_comp
        
        if use_ifrs_auto:
            try:
                parts = str(periodo_actual).split('-')
                if len(parts) >= 2:
                    year = int(parts[0])
                    month = parts[1]
                    prev_year = str(year - 1)
                    periodo_comp_balance = f"{prev_year}-12"
                    periodo_comp_resultados = f"{prev_year}-{month}"
            except Exception:
                pass
        
        from src.core.excel_utils import format_periodo
        txt_periodo_act = format_periodo(periodo_actual)
        txt_periodo_comp = format_periodo(periodo_comp) if periodo_comp != "Ninguno" else ""
        
        # Obtener utilidad neta
        val_utilidad_pesos = 0.0
        if is_consolidated:
            grupo_name = empresa_activa.replace("[GRUPO] ", "").strip()
            db = SessionLocal()
            try:
                grupo_obj = db.query(ConsolidationGroup).filter_by(nombre_grupo=grupo_name).first()
                grupo_id = grupo_obj.id if grupo_obj else None
            finally:
                db.close()
                
            if grupo_id:
                from src.core.consolidacion_engine import generar_hoja_trabajo
                df_hoja_act, _ = generar_hoja_trabajo(grupo_id, periodo_actual)
                if df_hoja_act is not None:
                    idx_er_act = df_hoja_act[df_hoja_act['Balance clasificado'] == "Estado de Resultados"].index
                    df_sec_act = df_hoja_act.loc[idx_er_act[0]+1:] if not idx_er_act.empty else df_hoja_act
                    row_hoja_act = df_sec_act[df_sec_act['Balance clasificado'].astype(str).str.strip() == "Ganancias (Pérdida) del Ejercicio"]
                    val_utilidad_pesos = -1 * float(row_hoja_act['CONSOLIDADO'].values[0]) if not row_hoja_act.empty else 0.0
        else:
            val_utilidad_pesos = -1 * PlCuboDB.get_pl_cubo_total_sum(empresa_activa, periodo_actual)
            
        txt_utilidad_neta = ""
        if val_utilidad_pesos < 0:
            txt_utilidad_neta = f"({abs(val_utilidad_pesos / scale_factor):,.0f})".replace(",", ".")
        else:
            txt_utilidad_neta = f"{val_utilidad_pesos / scale_factor:,.0f}".replace(",", ".")
            
        # Contexto de variables Jinja2 para docxtpl
        context = {
            "EMPRESA_ACTIVA": empresa_activa.replace("[GRUPO] ", "").strip(),
            "PERIODO_ACTUAL": txt_periodo_act,
            "PERIODO_COMPARATIVO": txt_periodo_comp,
            "UTILIDAD_NETA": txt_utilidad_neta,
            "empresa_activa": empresa_activa.replace("[GRUPO] ", "").strip(),
            "periodo_actual": txt_periodo_act,
            "periodo_comparativo": txt_periodo_comp,
            "utilidad_neta": txt_utilidad_neta
        }

        # Ocultar temporalmente los marcadores programáticos {{# para evitar errores de sintaxis en Jinja2
        def mask_programmatic_tags(doc_obj):
            for p in doc_obj.paragraphs:
                if "{{#" in p.text:
                    p.text = p.text.replace("{{#", "__HASH_TAG__")
            for table in doc_obj.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if "{{#" in p.text:
                                p.text = p.text.replace("{{#", "__HASH_TAG__")

        # Restaurar marcadores programáticos {{#
        def unmask_programmatic_tags(doc_obj):
            doc_to_use = doc_obj.docx if hasattr(doc_obj, 'docx') and doc_obj.docx is not None else doc_obj
            for p in doc_to_use.paragraphs:
                if "__HASH_TAG__" in p.text:
                    p.text = p.text.replace("__HASH_TAG__", "{{#")
            for table in doc_to_use.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if "__HASH_TAG__" in p.text:
                                p.text = p.text.replace("__HASH_TAG__", "{{#")

        # Cargar plantilla con python-docx primero para enmascarar de forma segura
        try:
            self.template_bytes_io.seek(0)
            doc_mask = Document(self.template_bytes_io)
            mask_programmatic_tags(doc_mask)
            
            temp_io = BytesIO()
            doc_mask.save(temp_io)
            temp_io.seek(0)
            
            doc = DocxTemplate(temp_io)
        except Exception as e:
            # Fallback en caso de problemas con la carga inicial
            self.template_bytes_io.seek(0)
            doc = DocxTemplate(self.template_bytes_io)
            try:
                mask_programmatic_tags(doc)
            except Exception:
                pass

        # Renderizar variables Jinja2 de forma nativa con docxtpl
        doc.render(context)
        
        unmask_programmatic_tags(doc)
                        
        # 3. Procesar e inyectar reportes principales y notas
        paragraphs_to_process = list(doc.paragraphs)
        paragraphs_to_delete = []
        
        for p in paragraphs_to_process:
            tag = p.text.strip()
            if not (tag.startswith("{{#") and tag.endswith("}}")):
                continue
                
            tag_content = tag.replace("{{", "").replace("}}", "").strip() # ej. #BALANCE:2025-12 o #BALANCE o #N04:2025-06
            if ":" in tag_content:
                parts = tag_content.split(":", 1)
                code = parts[0].strip() # ej. #BALANCE
                override_period = parts[1].strip() # ej. 2025-12
            else:
                code = tag_content
                override_period = None
            
            # --- CASO REPORTES PRINCIPALES ---
            if code in ["#BALANCE", "#ER", "#EFE", "#ORI", "#PATRIMONIO"]:
                # Generar reporte a DataFrame
                preview_df = None
                
                # Seleccionar comparativo dinámico o manual
                comp_period = override_period if override_period else (periodo_comp_balance if code == "#BALANCE" else periodo_comp_resultados)
                
                if code == "#BALANCE":
                    if is_consolidated:
                        grupo_name = empresa_activa.replace("[GRUPO] ", "").strip()
                        db = SessionLocal()
                        try:
                            grupo_obj = db.query(ConsolidationGroup).filter_by(nombre_grupo=grupo_name).first()
                            grupo_id = grupo_obj.id
                        finally:
                            db.close()
                        from src.core.consolidacion_engine import generar_hoja_trabajo
                        df_hoja_act, _ = generar_hoja_trabajo(grupo_id, periodo_actual)
                        df_hoja_comp = None
                        if comp_period != "Ninguno":
                            df_hoja_comp, _ = generar_hoja_trabajo(grupo_id, comp_period)
                            
                        # Slice assets
                        idx_er = df_hoja_act[df_hoja_act['Balance clasificado'] == "Estado de Resultados"].index
                        df_hoja_act_sec = df_hoja_act.loc[:idx_er[0]-1] if not idx_er.empty else df_hoja_act
                        df_hoja_act_clean = df_hoja_act_sec[df_hoja_act_sec['Balance clasificado'].notna() & (df_hoja_act_sec['Balance clasificado'].str.strip() != "")]
                        
                        tb_df = pd.DataFrame({'cuenta_id': df_hoja_act_clean['Balance clasificado'], 'saldo_final': df_hoja_act_clean['CONSOLIDADO']})
                        tb_df_comp = None
                        if df_hoja_comp is not None:
                            idx_er_comp = df_hoja_comp[df_hoja_comp['Balance clasificado'] == "Estado de Resultados"].index
                            df_hoja_comp_sec = df_hoja_comp.loc[:idx_er_comp[0]-1] if not idx_er_comp.empty else df_hoja_comp
                            df_hoja_comp_clean = df_hoja_comp_sec[df_hoja_comp_sec['Balance clasificado'].notna() & (df_hoja_comp_sec['Balance clasificado'].str.strip() != "")]
                            tb_df_comp = pd.DataFrame({'cuenta_id': df_hoja_comp_clean['Balance clasificado'], 'saldo_final': df_hoja_comp_clean['CONSOLIDADO']})
                            
                        dummy_map = pd.DataFrame({'N° de Cuenta': tb_df['cuenta_id'], 'Clasificación balance': tb_df['cuenta_id']})
                        from src.reporting.balance_generator import BalanceGenerator
                        gen = BalanceGenerator(os.path.join(empresa_path, "Balance clasificado.xlsx"))
                        result_bytes = gen.generate(tb_df, dummy_map, scale_factor=scale_factor, tb_df_comp=tb_df_comp, periodo_actual_str=periodo_actual, periodo_comp_str=comp_period if comp_period != "Ninguno" else None)
                        col_actual = str(periodo_actual)
                        col_comp = str(comp_period) if comp_period != "Ninguno" else "Comp"
                        preview_df = self._extract_balance_preview(result_bytes, col_actual, col_comp)
                    else:
                        tb_df = TrialBalanceDB.get_trial_balance(empresa_activa, periodo_actual)
                        tb_comp = TrialBalanceDB.get_trial_balance(empresa_activa, comp_period) if comp_period != "Ninguno" else None
                        from src.reporting.balance_generator import BalanceGenerator
                        gen = BalanceGenerator(os.path.join(empresa_path, "Balance clasificado.xlsx"))
                        result_bytes = gen.generate(tb_df, map_balance_df, scale_factor=scale_factor, tb_df_comp=tb_comp, periodo_actual_str=periodo_actual, periodo_comp_str=comp_period if comp_period != "Ninguno" else None)
                        col_actual = str(periodo_actual)
                        col_comp = str(comp_period) if comp_period != "Ninguno" else "Comp"
                        preview_df = self._extract_balance_preview(result_bytes, col_actual, col_comp)
                        
                elif code == "#ER":
                    if is_consolidated:
                        grupo_name = empresa_activa.replace("[GRUPO] ", "").strip()
                        db = SessionLocal()
                        try:
                            grupo_obj = db.query(ConsolidationGroup).filter_by(nombre_grupo=grupo_name).first()
                            grupo_id = grupo_obj.id
                        finally:
                            db.close()
                        from src.core.consolidacion_engine import generar_hoja_trabajo
                        df_hoja_act, _ = generar_hoja_trabajo(grupo_id, periodo_actual)
                        df_hoja_comp = None
                        if comp_period != "Ninguno":
                            df_hoja_comp, _ = generar_hoja_trabajo(grupo_id, comp_period)
                            
                        # Slice PL
                        idx_er = df_hoja_act[df_hoja_act['Balance clasificado'] == "Estado de Resultados"].index
                        df_hoja_act_sec = df_hoja_act.loc[idx_er[0]+1:] if not idx_er.empty else df_hoja_act
                        df_hoja_act_clean = df_hoja_act_sec[df_hoja_act_sec['Balance clasificado'].notna() & (df_hoja_act_sec['Balance clasificado'].str.strip() != "")]
                        pl_dict_act = {row['Balance clasificado']: [row['CONSOLIDADO']] for _, row in df_hoja_act_clean.iterrows()}
                        pl_df_wide = pd.DataFrame(pl_dict_act)
                        
                        pl_df_comp_wide = None
                        if df_hoja_comp is not None:
                            idx_er_comp = df_hoja_comp[df_hoja_comp['Balance clasificado'] == "Estado de Resultados"].index
                            df_hoja_comp_sec = df_hoja_comp.loc[idx_er_comp[0]+1:] if not idx_er_comp.empty else df_hoja_comp
                            df_hoja_comp_clean = df_hoja_comp_sec[df_hoja_comp_sec['Balance clasificado'].notna() & (df_hoja_comp_sec['Balance clasificado'].str.strip() != "")]
                            pl_dict_comp = {row['Balance clasificado']: [row['CONSOLIDADO']] for _, row in df_hoja_comp_clean.iterrows()}
                            pl_df_comp_wide = pd.DataFrame(pl_dict_comp)
                            
                        from src.reporting.er_generator import ERGenerator
                        gen = ERGenerator(os.path.join(empresa_path, "Estado de Resultados Clasificados.xlsx"))
                        result_bytes, _ = gen.generate(pl_df_wide, scale_factor=scale_factor, pl_df_comp=pl_df_comp_wide, periodo_actual_str=periodo_actual, periodo_comp_str=comp_period if comp_period != "Ninguno" else None)
                        col_actual = str(periodo_actual)
                        col_comp = str(comp_period) if comp_period != "Ninguno" else "Comp"
                        preview_df = self._extract_er_preview(result_bytes, col_actual, col_comp)
                    else:
                        pl_df_wide = PlCuboDB.get_pl_cubo(empresa_activa, periodo_actual)
                        pl_df_comp_wide = PlCuboDB.get_pl_cubo(empresa_activa, comp_period) if comp_period != "Ninguno" else None
                        
                        from src.reporting.er_generator import ERGenerator
                        gen = ERGenerator(os.path.join(empresa_path, "Estado de Resultados Clasificados.xlsx"))
                        result_bytes, _ = gen.generate(pl_df_wide, scale_factor=scale_factor, pl_df_comp=pl_df_comp_wide, periodo_actual_str=periodo_actual, periodo_comp_str=comp_period if comp_period != "Ninguno" else None)
                        col_actual = str(periodo_actual)
                        col_comp = str(comp_period) if comp_period != "Ninguno" else "Comp"
                        preview_df = self._extract_er_preview(result_bytes, col_actual, col_comp)
                        
                elif code == "#ORI":
                    if is_consolidated:
                        grupo_name = empresa_activa.replace("[GRUPO] ", "").strip()
                        db = SessionLocal()
                        try:
                            grupo_obj = db.query(ConsolidationGroup).filter_by(nombre_grupo=grupo_name).first()
                            grupo_id = grupo_obj.id
                        finally:
                            db.close()
                        from src.core.consolidacion_engine import generar_hoja_trabajo
                        df_hoja_act, _ = generar_hoja_trabajo(grupo_id, periodo_actual)
                        df_hoja_comp = None
                        if comp_period != "Ninguno":
                            df_hoja_comp, _ = generar_hoja_trabajo(grupo_id, comp_period)
                            
                        # Slice PL
                        idx_er = df_hoja_act[df_hoja_act['Balance clasificado'] == "Estado de Resultados"].index
                        df_hoja_act_sec = df_hoja_act.loc[idx_er[0]+1:] if not idx_er.empty else df_hoja_act
                        df_hoja_act_clean = df_hoja_act_sec[df_hoja_act_sec['Balance clasificado'].notna() & (df_hoja_act_sec['Balance clasificado'].str.strip() != "")]
                        pl_dict_act = {row['Balance clasificado']: [row['CONSOLIDADO']] for _, row in df_hoja_act_clean.iterrows()}
                        pl_df_wide = pd.DataFrame(pl_dict_act)
                        
                        pl_df_comp_wide = None
                        if df_hoja_comp is not None:
                            idx_er_comp = df_hoja_comp[df_hoja_comp['Balance clasificado'] == "Estado de Resultados"].index
                            df_hoja_comp_sec = df_hoja_comp.loc[idx_er_comp[0]+1:] if not idx_er_comp.empty else df_hoja_comp
                            df_hoja_comp_clean = df_hoja_comp_sec[df_hoja_comp_sec['Balance clasificado'].notna() & (df_hoja_comp_sec['Balance clasificado'].str.strip() != "")]
                            pl_dict_comp = {row['Balance clasificado']: [row['CONSOLIDADO']] for _, row in df_hoja_comp_clean.iterrows()}
                            pl_df_comp_wide = pd.DataFrame(pl_dict_comp)
                            
                        from src.reporting.er_generator import ERGenerator
                        gen_er = ERGenerator(os.path.join(empresa_path, "Estado de Resultados Clasificados.xlsx"))
                        _, preview_er = gen_er.generate(pl_df_wide, scale_factor=scale_factor, pl_df_comp=pl_df_comp_wide, periodo_actual_str=periodo_actual, periodo_comp_str=comp_period if comp_period != "Ninguno" else None)
                    else:
                        pl_df_wide = PlCuboDB.get_pl_cubo(empresa_activa, periodo_actual)
                        pl_df_comp_wide = PlCuboDB.get_pl_cubo(empresa_activa, comp_period) if comp_period != "Ninguno" else None
                        
                        from src.reporting.er_generator import ERGenerator
                        gen_er = ERGenerator(os.path.join(empresa_path, "Estado de Resultados Clasificados.xlsx"))
                        _, preview_er = gen_er.generate(pl_df_wide, scale_factor=scale_factor, pl_df_comp=pl_df_comp_wide, periodo_actual_str=periodo_actual, periodo_comp_str=comp_period if comp_period != "Ninguno" else None)
                    
                    from src.reporting.ori_generator import OriGenerator
                    template_ori_path = os.path.join(empresa_path, "Estado de Resultados Integrales.xlsx")
                    if os.path.exists(template_ori_path):
                        gen_ori = OriGenerator(template_ori_path)
                        result_bytes = gen_ori.generate(
                            preview_er, 
                            periodo_actual_str=str(periodo_actual), 
                            periodo_comp_str=str(comp_period) if comp_period != "Ninguno" else None,
                            empresa=empresa_activa
                        )
                        col_actual = str(periodo_actual)
                        col_comp = str(comp_period) if comp_period != "Ninguno" else "Comp"
                        preview_df = self._extract_ori_preview(result_bytes, col_actual, col_comp)
                    else:
                        preview_df = None
                        
                elif code == "#PATRIMONIO":
                    template_pat_path = os.path.join(empresa_path, "Estado de Cambios en el Patrimonio.xlsx")
                    if os.path.exists(template_pat_path):
                        from src.reporting.patrimonio_generator import PatrimonioGenerator
                        pat_engine = PatrimonioGenerator(template_pat_path)
                        ex_pat = pat_engine.generate(
                            bal_preview_df=None,
                            pl_preview_df=None,
                            periodo_actual_str=str(periodo_actual),
                            periodo_comp_str=str(comp_period) if comp_period != "Ninguno" else None,
                            empresa=empresa_activa
                        )
                        ex_pat.seek(0)
                        from src.core.excel_utils import detect_patrimonio_skiprows, clean_preview_dataframe
                        pat_skip = detect_patrimonio_skiprows(ex_pat)
                        ex_pat.seek(0)
                        preview_df = pd.read_excel(ex_pat, skiprows=pat_skip)
                        preview_df = clean_preview_dataframe(preview_df)
                        first_col = preview_df.columns[0]
                        preview_df = preview_df.dropna(how='all', subset=[first_col])
                    else:
                        preview_df = None
                        
                elif code == "#EFE":
                    cf_filename = "Estado de Flujos de Efectivo.xlsx"
                    template_cf_path = os.path.join(empresa_path, cf_filename)
                    if not os.path.exists(template_cf_path):
                        import shutil
                        shutil.copy2(os.path.join("templates", cf_filename), template_cf_path)
                        
                    from src.reporting.cash_flow_generator import CashFlowGenerator
                    gen = CashFlowGenerator(template_cf_path)
                    
                    if is_consolidated:
                        from src.core.consolidacion_engine import generar_hoja_trabajo
                        db = SessionLocal()
                        try:
                            grupo_obj = db.query(ConsolidationGroup).filter_by(nombre_grupo=empresa_activa.replace("[GRUPO] ", "").strip()).first()
                            grupo_id = grupo_obj.id if grupo_obj else None
                        finally:
                            db.close()
                        df_hoja_act, _ = generar_hoja_trabajo(grupo_id, periodo_actual)
                        df_hoja_comp = None
                        if comp_period != "Ninguno":
                            df_hoja_comp, _ = generar_hoja_trabajo(grupo_id, comp_period)
                            
                        result_bytes, _ = gen.generate(
                            empresa=empresa_activa,
                            periodo_actual_str=periodo_actual,
                            periodo_comp_str=comp_period if comp_period != "Ninguno" else None,
                            map_balance_df=map_balance_df,
                            map_pl_df=map_pl_df,
                            method="Directo",
                            is_consolidado=True,
                            consolidated_hoja_trabajo_df=df_hoja_act,
                            consolidated_hoja_trabajo_comp_df=df_hoja_comp,
                            scale_factor=scale_factor
                        )
                    else:
                        result_bytes, _ = gen.generate(
                            empresa=empresa_activa,
                            periodo_actual_str=periodo_actual,
                            periodo_comp_str=comp_period if comp_period != "Ninguno" else None,
                            map_balance_df=map_balance_df,
                            map_pl_df=map_pl_df,
                            method="Directo",
                            is_consolidado=False,
                            scale_factor=scale_factor
                        )
                        
                    col_actual = str(periodo_actual)
                    col_comp = str(comp_period) if comp_period != "Ninguno" else "Comp"
                    preview_df = self._extract_cf_preview(result_bytes, col_actual, col_comp)
                    
                if preview_df is not None:
                    # Inyectar tabla en la posición de p
                    tbl = doc.add_table(rows=len(preview_df) + 1, cols=len(preview_df.columns))
                    numeric_cols = preview_df.select_dtypes(include=['number']).columns.tolist()
                    populate_docx_table(tbl, preview_df, numeric_cols)
                    
                    move_element_before(p, tbl._tbl)
                    paragraphs_to_delete.append(p)
                    
            # --- CASO NOTAS ---
            else:
                base_code = code
                table_idx = None
                if "." in code:
                    parts_dot = code.split(".", 1)
                    base_code = parts_dot[0].strip()
                    try:
                        table_idx = int(parts_dot[1].strip()) - 1
                    except ValueError:
                        pass

                from src.reporting.notes import NOTE_REGISTRY
                if base_code in NOTE_REGISTRY:
                    info = NOTE_REGISTRY[base_code]
                    target_sheets = info['sheets']
                    
                    if target_sheets:
                        template_nota = "Plantilla de notas_v1.xlsx"
                        if os.path.exists(template_nota):
                            category = info.get('category', 'resultados')
                            is_balance_note = category in ["activos_corrientes", "activos_no_corrientes", "pasivos_corrientes", "pasivos_no_corrientes", "patrimonio"]
                            note_comp = override_period if override_period else (periodo_comp_balance if is_balance_note else periodo_comp_resultados)
                            
                            from src.ui_pages.informes_y_notas import load_all_entity_contexts
                            entity_contexts = load_all_entity_contexts(
                                active_entity=empresa_activa,
                                periodo_actual=periodo_actual,
                                periodo_comp=note_comp,
                                map_balance_df=map_balance_df,
                                map_pl_df=map_pl_df
                            )
                            
                            from src.reporting.note_generator import NoteGenerator
                            engine = NoteGenerator(template_nota)
                            excel_nota_out = engine.generate(
                                sheet_names=target_sheets,
                                entity_contexts=entity_contexts,
                                active_entity_name=empresa_activa,
                                is_consolidated=is_consolidated,
                                scale_factor=scale_factor,
                                periodo_actual_str=periodo_actual,
                                periodo_comp_str=note_comp,
                                map_balance_df=map_balance_df,
                                map_pl_df=map_pl_df
                            )
                            
                            # Evaluar las fórmulas antes de procesar para evitar NaNs en totales
                            from src.ui_pages.informes_y_notas import evaluate_formulas_in_workbook
                            excel_nota_out = evaluate_formulas_in_workbook(excel_nota_out)
                            
                            # Intentar extraer usando rangos nombrados
                            from src.core.excel_utils import extract_named_ranges_from_excel
                            named_ranges = extract_named_ranges_from_excel(excel_nota_out, base_code)
                            
                            if named_ranges:
                                elements = []
                                for r_name, df_chunk in named_ranges:
                                    elements.append(("table", df_chunk))
                            else:
                                excel_nota_out.seek(0)
                                preview_nota_df = pd.read_excel(excel_nota_out, sheet_name=target_sheets[0], header=None)
                                from src.ui_pages.informes_y_notas import split_sheet_into_elements
                                elements = split_sheet_into_elements(preview_nota_df)
                            
                            # Si se solicitó un índice de tabla específico (ej: #N04.1)
                            if table_idx is not None:
                                tables = [el for el in elements if el[0] == "table"]
                                if 0 <= table_idx < len(tables):
                                    elements = [tables[table_idx]]
                                else:
                                    elements = [] # Fuera de rango o sin tablas
                            
                            # Inyectar elementos uno a uno antes del párrafo p
                            table_counter = 0
                            for el_type, el_val in elements:
                                if el_type == "text":
                                    new_p = doc.add_paragraph()
                                    run = new_p.add_run(str(el_val))
                                    run.font.bold = True
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(9)
                                    move_element_before(p, new_p._p)
                                else:
                                    chunk_df = el_val.dropna(how='all', axis=0).reset_index(drop=True)
                                    if chunk_df.empty:
                                        continue
                                    
                                    table_counter += 1
                                    tbl = doc.add_table(rows=len(chunk_df), cols=len(chunk_df.columns))
                                    populate_notes_docx_table(tbl, chunk_df)
                                    
                                    move_element_before(p, tbl._tbl)
                                    
                            paragraphs_to_delete.append(p)
                            
        # Eliminar párrafos marcadores procesados
        parent = doc.element.body
        for p in paragraphs_to_delete:
            try:
                parent.remove(p._p)
            except:
                pass
                
        # Guardar en buffer
        output_buffer = BytesIO()
        doc.save(output_buffer)
        output_buffer.seek(0)
        return output_buffer
