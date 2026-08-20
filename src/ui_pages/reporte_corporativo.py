import streamlit as st
import pandas as pd
import os
from io import BytesIO
from src.core.excel_utils import format_periodo
from src.core.word_template_engine import WordTemplateEngine

def generate_sample_template():
    """Genera en memoria una plantilla Word de ejemplo con las etiquetas soportadas."""
    from docx import Document
    doc = Document()
    doc.add_heading("INFORME FINANCIERO CORPORATIVO", level=0)
    
    p = doc.add_paragraph()
    p.add_run("Este documento ha sido generado automáticamente utilizando la plantilla de ejemplo y el motor de reportes de ").italic = True
    p.add_run("Informes Pro").bold = True
    p.add_run(".")
    
    doc.add_heading("1. Resumen General", level=1)
    doc.add_paragraph("Empresa Activa: {{EMPRESA_ACTIVA}}")
    doc.add_paragraph("Período de Análisis YTD: {{PERIODO_ACTUAL}}")
    doc.add_paragraph("Período Comparativo: {{PERIODO_COMPARATIVO}}")
    doc.add_paragraph("Utilidad Neta de las Operaciones: {{UTILIDAD_NETA}}")
    
    doc.add_heading("2. Estado de Situación Financiera (Balance)", level=1)
    doc.add_paragraph("El siguiente cuadro detalla la situación financiera consolidada/individual de la compañía:")
    doc.add_paragraph("{{#BALANCE}}")
    
    doc.add_heading("3. Estado de Resultados (P&L)", level=1)
    doc.add_paragraph("Desempeño de ingresos, costos y gastos acumulados para el ejercicio:")
    doc.add_paragraph("{{#ER}}")
    
    doc.add_heading("4. Estado de Flujos de Efectivo (EFE)", level=1)
    doc.add_paragraph("Origen y aplicación de recursos de efectivo:")
    doc.add_paragraph("{{#EFE}}")
    
    doc.add_heading("5. Notas Explicativas a los Estados Financieros", level=1)
    
    doc.add_heading("Nota 4: Efectivo y Equivalentes de Efectivo", level=2)
    doc.add_paragraph("{{#N04}}")
    
    doc.add_heading("Nota 14: Partes Relacionadas", level=2)
    doc.add_paragraph("{{#N14}}")
    
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def render(empresa_seleccionada, empresa_path):
    if empresa_seleccionada == "[GLOBAL] Configuración General" or "GLOBAL" in empresa_seleccionada:
        st.warning("Módulo de Sociedad Activa: Por favor, selecciona una empresa de trabajo específica (ej. Pacifico SpA) en la barra lateral izquierda para acceder a esta sección.")
        st.stop()
        
    # Detectar cambio de empresa activa y resetear la plantilla del reporte anterior
    if st.session_state.get('corp_empresa_activa') != empresa_seleccionada:
        for k in ['corp_word_output_bytes', 'corp_word_filename', 'rep_corp_file']:
            if k in st.session_state:
                del st.session_state[k]
        st.session_state['corp_empresa_activa'] = empresa_seleccionada
        
    st.title("Generación de reportes Word")
    st.markdown("""
    Genera informes corporativos y de gestión a partir de **plantillas personalizadas de Word (`.docx`)**. 
    El motor inyectará automáticamente los estados financieros consolidados o individuales y las notas explicativas en las marcas correspondientes.
    """)
    
    # Acordeón de Guía de Etiquetas
    with st.expander("Guía de Diseño de Plantilla (Etiquetas Soportadas)", expanded=False):
        st.markdown("""
        Puedes diseñar tu plantilla de Word en cualquier formato, incluir portadas, logotipos corporativos, 
        encabezados y tablas estilizadas. Para inyectar datos dinámicos, utiliza las siguientes etiquetas en el texto:
        
        #### Variables Simples (Párrafos o Celdas):
        * `{{EMPRESA_ACTIVA}}` : Nombre de la empresa o grupo seleccionado.
        * `{{PERIODO_ACTUAL}}` : Período actual formateado (ej. *31 de Diciembre de 2025*).
        * `{{PERIODO_COMPARATIVO}}` : Período comparativo formateado (ej. *31 de Diciembre de 2024* o vacío si no se selecciona).
        * `{{UTILIDAD_NETA}}` : Utilidad neta formateada con el factor de escala seleccionado.
        
        #### Tablas de Reportes Principales:
        Inserta la etiqueta en una línea sola para inyectar la tabla completa:
        * `{{#BALANCE}}` : Inyecta el Estado de Situación Financiera.
        * `{{#ER}}` : Inyecta el Estado de Resultados.
        * `{{#EFE}}` : Inyecta el Estado de Flujos de Efectivo.
        * `{{#PATRIMONIO}}` : Inyecta el Estado de Cambios en el Patrimonio.
        * `{{#ORI}}` : Inyecta el Estado de Resultados Integrales.
        
        #### Notas Explicativas:
        Inserta la etiqueta de la nota para inyectar su respectiva explicación de texto y cuadros asociados:
        * `{{#N04}}` : Efectivo y Equivalentes.
        * `{{#N06}}` : Deudores Comerciales.
        * `{{#N14}}` : Partes Relacionadas.
        * *(Cualquier nota registrada con su respectivo prefijo `#NXX`)*
        """)
        
        # Ofrecer descarga de la plantilla de ejemplo
        sample_buf = generate_sample_template()
        st.download_button(
            label="📥 Descargar Plantilla de Ejemplo (.docx)",
            data=sample_buf,
            file_name="Plantilla_Ejemplo_Reporte.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="dl_sample_template"
        )
        
    st.write("")
    
    # 1. Obtener períodos históricos disponibles para los selectores
    from src.models.database import SessionLocal
    from src.models.historical_data import HistoricalDataRecord
    
    db = SessionLocal()
    try:
        per_recs = db.query(HistoricalDataRecord.periodo).distinct().all()
        periodos_hist = sorted([r[0] for r in per_recs], reverse=True)
    except Exception as e:
        periodos_hist = []
    finally:
        db.close()
        
    if not periodos_hist:
        periodos_hist = ["2025-12", "2025-03", "2024-12"]
        
    # 2. Configuración de parámetros en columnas
    st.subheader("⚙️ Configuración del Informe")
    col1, col2, col3 = st.columns([2, 2, 2])
    
    with col1:
        periodo_actual = st.selectbox(
            "Período Actual:",
            periodos_hist,
            index=0,
            format_func=format_periodo,
            key="rep_corp_act"
        )
    with col2:
        periodos_comp_opts = ["Ninguno"] + periodos_hist
        # Pre-seleccionar el segundo periodo si existe
        default_comp_idx = 1 if len(periodos_hist) > 1 else 0
        periodo_comp = st.selectbox(
            "Período Comparativo (Opcional):",
            periodos_comp_opts,
            index=default_comp_idx,
            format_func=lambda x: "Ninguno" if x == "Ninguno" else format_periodo(x),
            key="rep_corp_comp"
        )
    with col3:
        unidad = st.selectbox(
            "Unidad de Medida / Escala:",
            ["Miles de pesos (M$)", "Unidades ($)", "Millones de pesos (MM$)"],
            index=0,
            key="rep_corp_scale"
        )
        
        if "Miles" in unidad:
            scale_factor = 1000.0
        elif "Millones" in unidad:
            scale_factor = 1000000.0
        else:
            scale_factor = 1.0

    st.write("")
    
    # Checkbox para habilitar Modo Inteligente IFRS
    use_ifrs_auto = st.checkbox(
        "Reglas comparativas IFRS automáticas", 
        value=True, 
        help="Si está activo: el Balance comparará con el cierre de diciembre del año anterior, y el Estado de Resultados / Flujo de Efectivo / Notas compararán con el mismo mes del año anterior. Si se desactiva: se utilizará la opción del selector 'Período Comparativo' para todos los rubros.",
        key="rep_corp_ifrs_auto"
    )
    
    st.write("")
    
    # 3. Panel de Carga de Plantilla
    st.subheader("📂 Plantilla Word")
    
    def reset_corp_report_state():
        for k in ['corp_word_output_bytes', 'corp_word_filename', 'rep_corp_file']:
            if k in st.session_state:
                del st.session_state[k]

    uploaded_file = st.file_uploader(
        "Carga tu archivo de plantilla Word (.docx):",
        type=["docx"],
        key="rep_corp_file"
    )
    
    if uploaded_file is not None:
        st.success("Plantilla cargada con éxito. Listo para generar el reporte.")
        
        # Botón para procesar
        if st.button("Generar Reporte Corporativo", type="primary", use_container_width=True):
            with st.spinner("Procesando plantilla e inyectando datos financieros..."):
                try:
                    # Crear una copia de los bytes subidos
                    template_bytes = BytesIO(uploaded_file.read())
                    
                    # Instanciar el motor y ejecutar
                    engine = WordTemplateEngine(template_bytes)
                    output_buffer = engine.process(
                        empresa_activa=empresa_seleccionada,
                        periodo_actual=periodo_actual,
                        periodo_comp=periodo_comp,
                        scale_factor=scale_factor,
                        use_ifrs_auto=use_ifrs_auto
                    )
                    
                    out_bytes = output_buffer.getvalue() if hasattr(output_buffer, 'getvalue') else output_buffer
                    st.session_state['corp_word_output_bytes'] = out_bytes
                    st.session_state['corp_word_filename'] = f"Reporte_Corporativo_{empresa_seleccionada.replace('[GRUPO] ', '').replace(' ', '_')}_{periodo_actual}.docx"
                    st.success("¡Reporte generado con éxito!")
                except Exception as e:
                    st.error(f"Ocurrió un error al procesar el reporte: {str(e)}")
                    st.exception(e)
                    
        if 'corp_word_output_bytes' in st.session_state:
            st.write("---")
            col_dl1, col_dl2 = st.columns([3, 1])
            with col_dl1:
                st.download_button(
                    label="Descargar Reporte Final (.docx)",
                    data=st.session_state['corp_word_output_bytes'],
                    file_name=st.session_state.get('corp_word_filename', 'Reporte_Corporativo.docx'),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                    use_container_width=True,
                    on_click=reset_corp_report_state,
                    key="dl_processed_report"
                )
            with col_dl2:
                if st.button("Limpiar / Listo para otro reporte", use_container_width=True, key="btn_reset_corp"):
                    reset_corp_report_state()
                    st.rerun()
    else:
        st.info("Sube una plantilla en formato Word (.docx) o descarga nuestra plantilla de ejemplo para comenzar a generar reportes.")
