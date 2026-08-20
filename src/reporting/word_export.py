from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from io import BytesIO
import pandas as pd

def set_cell_bg(cell, fill_color):
    """Inyecta XML OXML para dar color de fondo a una celda de Word."""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill_color))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=50, bottom=50, left=100, right=100):
    """Ajusta los márgenes internos (padding) de la celda de Word en dxa (1 pt = 20 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_cell_borders(cell, top=None, bottom=None):
    """Inyecta bordes OXML a una celda de Word (útil para filas de totales/subtotales contables)."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders_xml = f'<w:tcBorders {nsdecls("w")}>'
    if top:
        borders_xml += f'<w:top w:val="{top.get("val", "single")}" w:sz="{top.get("sz", "6")}" w:space="0" w:color="{top.get("color", "000000")}"/>'
    else:
        borders_xml += f'<w:top w:val="none"/>'
    if bottom:
        borders_xml += f'<w:bottom w:val="{bottom.get("val", "single")}" w:sz="{bottom.get("sz", "6")}" w:space="0" w:color="{bottom.get("color", "000000")}"/>'
    else:
        borders_xml += f'<w:bottom w:val="none"/>'
    borders_xml += f'<w:left w:val="none"/><w:right w:val="none"/>'
    borders_xml += '</w:tcBorders>'
    tcPr.append(parse_xml(borders_xml))

def apply_table_column_widths(table, num_cols, total_width_inches=6.8):
    """Calcula y aplica anchos de columna óptimos para informes financieros."""
    table.autofit = False
    
    if num_cols <= 1:
        col_widths = [total_width_inches]
    elif num_cols == 2:
        col_widths = [total_width_inches * 0.70, total_width_inches * 0.30]
    elif num_cols == 3:
        col_widths = [total_width_inches * 0.64, total_width_inches * 0.18, total_width_inches * 0.18]
    elif num_cols == 4:
        col_widths = [total_width_inches * 0.58, total_width_inches * 0.10, total_width_inches * 0.16, total_width_inches * 0.16]
    else:
        col0 = max(3.2, total_width_inches - (1.1 * (num_cols - 1)))
        rem = max(0.8, (total_width_inches - col0) / (num_cols - 1))
        col_widths = [col0] + [rem] * (num_cols - 1)

    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
        for j, cell in enumerate(row.cells):
            if j < len(col_widths):
                cell.width = Inches(col_widths[j])

    if len(table.rows) > 0:
        hdr_trPr = table.rows[0]._tr.get_or_add_trPr()
        hdr_trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))


class WordExportEngine:
    @staticmethod
    def generate_classified_balance_word(df, title="Estado de Situación Financiera Clasificado", unit="Ch$"):
        """
        Genera un informe en Word con formato financiero profesional tipo IFRS / Memoria Anual.
        Aplica anchos de columna proporcionados, interlíneas compactas y bordes contables.
        """
        doc = Document()
        
        # Establecer márgenes de página limpios (0.8 pulgadas / ~2cm)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.85)
            section.right_margin = Inches(0.85)

        # Encabezado principal
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in heading.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(16)
            r.font.color.rgb = RGBColor(31, 78, 120)
            r.font.bold = True
        
        # Subtítulo / Unidad de medida
        subtitle = doc.add_paragraph(f"Expresado en {unit}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in subtitle.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(9.5)
            r.font.italic = True
            r.font.color.rgb = RGBColor(80, 80, 80)
        
        p_space = doc.add_paragraph()
        p_space.paragraph_format.space_before = Pt(0)
        p_space.paragraph_format.space_after = Pt(4)
        
        num_cols = len(df.columns)
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = 'Normal Table'

        # Encabezados de tabla
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(df.columns):
            hdr_cells[i].text = str(col_name)
            set_cell_bg(hdr_cells[i], "1F4E78") # Azul Corporativo
            set_cell_margins(hdr_cells[i], top=70, bottom=70, left=100, right=100)
            
            for p in hdr_cells[i].paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.0
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.RIGHT
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    run.font.name = 'Arial' 
                    run.font.size = Pt(9)
        
        numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
        
        # Filas de datos
        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            first_cell_str = str(row.iloc[0]).strip()
            first_cell_clean = first_cell_str.lower()
            
            # Clasificación de tipo de fila contable
            is_grand_total = any(k in first_cell_clean for k in ["total activos", "total patrimonio y pasivos", "patrimonio total", "ganancia (pérdida) del ejercicio", "resultado del ejercicio"])
            is_subtotal = not is_grand_total and any(k in first_cell_clean for k in ["total", "sub total", "saldo final", "totales"])
            is_category_header = not is_subtotal and not is_grand_total and (
                first_cell_str in ["Activos", "Activos corrientes", "Activos no corrientes", "Patrimonio y pasivos", "Pasivos corrientes", "Pasivos no corrientes", "Patrimonio"] or
                all(pd.isna(row.iloc[col_i]) or str(row.iloc[col_i]).strip() == "" for col_i in range(1, num_cols))
            )

            # Colores de fondo
            if is_grand_total:
                bg_color = "EBF2FA"
            elif is_subtotal:
                bg_color = "F4F7FA"
            elif is_category_header:
                bg_color = "F9FAFB"
            else:
                is_even = (index % 2 == 0)
                bg_color = "FFFFFF" if is_even else "F9FAFB"

            for i, col_name in enumerate(df.columns):
                val = row[col_name]
                if pd.isna(val) or val == "None" or str(val) == "nan":
                    text_val = ""
                else:
                    if isinstance(val, (int, float)):
                        try:
                            if val < 0:
                                text_val = f"({abs(val):,.0f})".replace(",", ".")
                            else:
                                text_val = f"{val:,.0f}".replace(",", ".")
                        except:
                            text_val = str(val)
                    else:
                        text_val = str(val)
                
                cell = row_cells[i]
                cell.text = text_val
                set_cell_bg(cell, bg_color)
                set_cell_margins(cell, top=40, bottom=40, left=100, right=100)
                
                # Aplicar bordes contables
                if is_grand_total:
                    set_cell_borders(cell, top={'val': 'single', 'sz': '6'}, bottom={'val': 'double', 'sz': '12'})
                elif is_subtotal:
                    set_cell_borders(cell, top={'val': 'single', 'sz': '6'}, bottom={'val': 'single', 'sz': '6'})
                
                for p in cell.paragraphs:
                    p.paragraph_format.space_before = Pt(1.5)
                    p.paragraph_format.space_after = Pt(1.5)
                    p.paragraph_format.line_spacing = 1.0
                    
                    if col_name in numeric_cols or i > 0 or isinstance(val, (int, float)):
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                    for run in p.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(8.5)
                        if is_grand_total or is_subtotal or is_category_header:
                            run.font.bold = True
                            if is_grand_total or is_category_header:
                                run.font.color.rgb = RGBColor(15, 45, 80)

        # Aplicar anchos de columna profesionales
        apply_table_column_widths(table, num_cols, total_width_inches=6.8)

        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    @staticmethod
    def generate_notes_word(elements, title="Nota", unit="M$", note_code=None):
        """
        Genera documento Word de Notas con formato contable ajustado,
        ancho de columnas proporcional e interlíneas compactas.
        """
        doc = Document()
        
        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.85)
            section.right_margin = Inches(0.85)
        
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in heading.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(15)
            r.font.color.rgb = RGBColor(31, 78, 120)
            r.font.bold = True
            
        subtitle = doc.add_paragraph(f"Expresado en {unit}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in subtitle.runs:
            r.font.name = 'Arial'
            r.font.size = Pt(9)
            r.font.italic = True
            
        doc.add_paragraph()
        
        table_counter = 0
        for el_type, el_val in elements:
            if el_type == "text":
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(str(el_val))
                run.font.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(9.5)
            else:
                chunk_df = el_val.dropna(how='all', axis=0).reset_index(drop=True)
                if chunk_df.empty:
                    continue
                    
                table_counter += 1
                sub_code = f"{note_code}.{table_counter}" if note_code else f"Cuadro {table_counter}"
                
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(f"[{sub_code}]")
                run.font.bold = True
                run.font.name = 'Arial'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(31, 78, 120)
                
                num_cols = len(chunk_df.columns)
                table = doc.add_table(rows=len(chunk_df), cols=num_cols)
                table.style = 'Normal Table'
                
                numeric_cols = chunk_df.select_dtypes(include=['number']).columns.tolist()
                
                in_header = True
                for row_idx, row in chunk_df.iterrows():
                    row_cells = table.rows[row_idx].cells
                    
                    has_data_numbers = False
                    for col_idx, val in enumerate(row):
                        if pd.notna(val) and isinstance(val, (int, float)):
                            if not (val > 2000 and val < 2100):
                                has_data_numbers = True
                                break
                    if has_data_numbers:
                        in_header = False
                        
                    first_cell_val = str(row.iloc[0]).lower().strip()
                    is_grand_total = any(x in first_cell_val for x in ["total general", "total final", "saldo final", "total patrimonio", "total activos"])
                    is_subtotal = not is_grand_total and any(x in first_cell_val for x in ["total", "sub total", "subtotal", "ganancia bruta"])
                    
                    if in_header:
                        bg_color = "1F4E78"
                    elif is_grand_total:
                        bg_color = "EBF2FA"
                    elif is_subtotal:
                        bg_color = "F4F7FA"
                    else:
                        is_even = (row_idx % 2 == 0)
                        bg_color = "FFFFFF" if is_even else "F9FAFB"
                    
                    for col_idx, col_name in enumerate(chunk_df.columns):
                        val = row[col_name]
                        if pd.isna(val) or val == "None" or str(val) == "nan" or str(val).strip() == "":
                            text_val = ""
                        else:
                            if isinstance(val, (int, float)):
                                try:
                                    if val < 0:
                                        text_val = f"({abs(val):,.0f})".replace(",", ".")
                                    else:
                                        text_val = f"{val:,.0f}".replace(",", ".")
                                except:
                                    text_val = str(val)
                            else:
                                text_val = str(val)
                                
                        cell = row_cells[col_idx]
                        cell.text = text_val
                        set_cell_bg(cell, bg_color)
                        set_cell_margins(cell, top=40, bottom=40, left=100, right=100)
                        
                        if is_grand_total:
                            set_cell_borders(cell, top={'val': 'single', 'sz': '6'}, bottom={'val': 'double', 'sz': '12'})
                        elif is_subtotal:
                            set_cell_borders(cell, top={'val': 'single', 'sz': '6'}, bottom={'val': 'single', 'sz': '6'})

                        for p in cell.paragraphs:
                            p.paragraph_format.space_before = Pt(1.5)
                            p.paragraph_format.space_after = Pt(1.5)
                            p.paragraph_format.line_spacing = 1.0
                            
                            if col_idx > 0 or col_idx in numeric_cols or isinstance(val, (int, float)):
                                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            else:
                                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                
                            for run in p.runs:
                                run.font.name = 'Arial'
                                run.font.size = Pt(8.5)
                                if in_header:
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(255, 255, 255)
                                elif is_grand_total or is_subtotal:
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(15, 45, 80)
                                    
                apply_table_column_widths(table, num_cols, total_width_inches=6.8)
                doc.add_paragraph()
                
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        return output


def generate_word_report(df, title="Reporte Financiero", subtitle="Expresado en pesos"):
    """
    Exporta un DataFrame genérico a Word con formato financiero profesional.
    Aplica para Estados de Resultados, Flujos de Efectivo, Patrimonio y Consolidados.
    """
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in heading.runs:
        r.font.name = 'Arial'
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(31, 78, 120)
        r.font.bold = True
    
    sub = doc.add_paragraph(subtitle)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.font.name = 'Arial'
        r.font.size = Pt(9.5)
        r.font.italic = True
        r.font.color.rgb = RGBColor(80, 80, 80)
        
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(0)
    p_space.paragraph_format.space_after = Pt(4)
    
    num_cols = len(df.columns)
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Normal Table'

    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(df.columns):
        hdr_cells[i].text = str(col_name)
        set_cell_bg(hdr_cells[i], "1F4E78") # Azul Corporativo
        set_cell_margins(hdr_cells[i], top=70, bottom=70, left=100, right=100)
        
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.name = 'Arial' 
                run.font.size = Pt(9)
                
    numeric_cols = df.select_dtypes(include=['number']).columns.tolist()
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        first_cell_str = str(row.iloc[0]).strip()
        first_cell_clean = first_cell_str.lower()
        
        is_grand_total = any(k in first_cell_clean for k in ["ganancia (pérdida) del ejercicio", "resultado del ejercicio", "total patrimonio", "total activos", "efectivo y equivalentes al efectivo al final del periodo"])
        is_subtotal = not is_grand_total and any(k in first_cell_clean for k in ["total", "sub total", "saldo final", "ganancia bruta", "resultado antes de", "flujos de efectivo procedentes de"])

        if is_grand_total:
            bg_color = "EBF2FA"
        elif is_subtotal:
            bg_color = "F4F7FA"
        else:
            is_even = (index % 2 == 0)
            bg_color = "FFFFFF" if is_even else "F9FAFB"

        for i, col_name in enumerate(df.columns):
            val = row[col_name]
            if pd.isna(val) or val == "None" or str(val) == "nan":
                text_val = ""
            else:
                if isinstance(val, (int, float)):
                    try:
                        if val < 0:
                            text_val = f"({abs(val):,.0f})".replace(",", ".")
                        else:
                            text_val = f"{val:,.0f}".replace(",", ".")
                    except:
                        text_val = str(val)
                else:
                    text_val = str(val)
            
            cell = row_cells[i]
            cell.text = text_val
            set_cell_bg(cell, bg_color)
            set_cell_margins(cell, top=40, bottom=40, left=100, right=100)
            
            if is_grand_total:
                set_cell_borders(cell, top={'val': 'single', 'sz': '6'}, bottom={'val': 'double', 'sz': '12'})
            elif is_subtotal:
                set_cell_borders(cell, top={'val': 'single', 'sz': '6'}, bottom={'val': 'single', 'sz': '6'})
            
            for p in row_cells[i].paragraphs:
                p.paragraph_format.space_before = Pt(1.5)
                p.paragraph_format.space_after = Pt(1.5)
                p.paragraph_format.line_spacing = 1.0
                
                if col_name in numeric_cols or isinstance(val, (int, float)) or i > 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(8.5)
                    if is_grand_total or is_subtotal:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(15, 45, 80)

    apply_table_column_widths(table, num_cols, total_width_inches=6.8)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()
